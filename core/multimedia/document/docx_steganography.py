"""
DOCX Steganography v3.0
디지털 포렌식 연구소

Microsoft Word 문서(DOCX) 기반 스테가노그래피 구현
DOCX의 XML 구조, 스타일, 속성을 조작하여 데이터를 은닉합니다.

주요 특징:
- 문서 속성 및 메타데이터 조작
- 텍스트 스타일 속성 미세 조정
- 공백 문자 및 탭 조작
- 숨김 텍스트 삽입
- 사용자 정의 XML 데이터 활용
"""

import os
import json
import hashlib
import struct
import zipfile
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Any, List, Tuple
from pathlib import Path
import logging
from Crypto.Cipher import AES
from Crypto.Random import get_random_bytes
from Crypto.Protocol.KDF import PBKDF2
import tempfile
import shutil
import base64

try:
    from python_docx import Document
    from python_docx.shared import Inches, Pt
    from python_docx.enum.text import WD_COLOR_INDEX
    from python_docx.oxml import OxmlElement
    from python_docx.oxml.ns import qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_COLOR_INDEX
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        PYTHON_DOCX_AVAILABLE = True
    except ImportError:
        Document = None
        Inches = None
        Pt = None
        WD_COLOR_INDEX = None
        OxmlElement = None
        qn = None
        PYTHON_DOCX_AVAILABLE = False

class DOCXSteganography:
    """DOCX 스테가노그래피 클래스"""
    
    def __init__(self,
                 metadata_method: bool = True,
                 style_method: bool = True,
                 whitespace_method: bool = True,
                 hidden_text_method: bool = True,
                 custom_xml_method: bool = True):
        """
        DOCX Steganography 초기화
        
        Args:
            metadata_method: 메타데이터 방식 사용 여부
            style_method: 스타일 조작 방식 사용 여부
            whitespace_method: 공백 문자 방식 사용 여부
            hidden_text_method: 숨김 텍스트 방식 사용 여부
            custom_xml_method: 사용자 정의 XML 방식 사용 여부
        """
        self.metadata_method = metadata_method
        self.style_method = style_method
        self.whitespace_method = whitespace_method
        self.hidden_text_method = hidden_text_method
        self.custom_xml_method = custom_xml_method
        
        # 로깅 설정
        self.logger = logging.getLogger(__name__)
        
        # 의존성 체크
        if not PYTHON_DOCX_AVAILABLE:
            self.logger.warning("python-docx를 찾을 수 없습니다. pip install python-docx로 설치하세요.")
        
        # DOCX 네임스페이스
        self.namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
            'dc': 'http://purl.org/dc/elements/1.1/',
            'dcterms': 'http://purl.org/dc/terms/',
            'dcmitype': 'http://purl.org/dc/dcmitype/',
            'xsi': 'http://www.w3.org/2001/XMLSchema-instance'
        }
        
        # 스타일 조작을 위한 미세한 변화값
        self.style_variations = {
            'font_size': [0.1, 0.2, 0.3, 0.4],  # 포인트 단위 미세 조정
            'spacing': [0.5, 1.0, 1.5, 2.0],    # 간격 조정
            'indent': [0.1, 0.2, 0.3, 0.4]      # 들여쓰기 조정
        }
    
    def _encrypt_data(self, data: str, password: str) -> bytes:
        """AES-256-GCM으로 데이터 암호화"""
        salt = get_random_bytes(16)
        key = PBKDF2(password, salt, 32, count=100000)
        
        cipher = AES.new(key, AES.MODE_GCM)
        ciphertext, auth_tag = cipher.encrypt_and_digest(data.encode('utf-8'))
        
        return salt + cipher.nonce + auth_tag + ciphertext
    
    def _decrypt_data(self, encrypted_data: bytes, password: str) -> str:
        """AES-256-GCM으로 데이터 복호화"""
        salt = encrypted_data[:16]
        nonce = encrypted_data[16:32]
        auth_tag = encrypted_data[32:48]
        ciphertext = encrypted_data[48:]
        
        key = PBKDF2(password, salt, 32, count=100000)
        cipher = AES.new(key, AES.MODE_GCM, nonce=nonce)
        
        return cipher.decrypt_and_verify(ciphertext, auth_tag).decode('utf-8')
    
    def _extract_docx_structure(self, docx_path: str) -> Dict[str, Any]:
        """DOCX 파일의 내부 구조 추출"""
        structure = {
            'document_xml': None,
            'core_properties': None,
            'app_properties': None,
            'custom_xml': [],
            'media_files': [],
            'relationships': None
        }
        
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # document.xml 추출
            try:
                with docx_zip.open('word/document.xml') as doc_file:
                    structure['document_xml'] = doc_file.read().decode('utf-8')
            except KeyError:
                pass
            
            # core properties 추출
            try:
                with docx_zip.open('docProps/core.xml') as core_file:
                    structure['core_properties'] = core_file.read().decode('utf-8')
            except KeyError:
                pass
            
            # app properties 추출
            try:
                with docx_zip.open('docProps/app.xml') as app_file:
                    structure['app_properties'] = app_file.read().decode('utf-8')
            except KeyError:
                pass
            
            # 사용자 정의 XML 파일들 찾기
            for file_info in docx_zip.infolist():
                if file_info.filename.startswith('customXml/'):
                    try:
                        with docx_zip.open(file_info.filename) as custom_file:
                            structure['custom_xml'].append({
                                'filename': file_info.filename,
                                'content': custom_file.read().decode('utf-8')
                            })
                    except:
                        pass
                
                # 미디어 파일들
                if file_info.filename.startswith('word/media/'):
                    structure['media_files'].append(file_info.filename)
        
        return structure
    
    def _embed_in_metadata(self, docx_path: str, data: bytes, output_path: str) -> bool:
        """메타데이터에 데이터 임베딩"""
        try:
            # 임시 디렉토리에서 작업
            with tempfile.TemporaryDirectory() as temp_dir:
                # DOCX 파일을 ZIP으로 추출
                with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                    docx_zip.extractall(temp_dir)
                
                # 코어 프로퍼티 수정
                core_path = os.path.join(temp_dir, 'docProps', 'core.xml')
                if os.path.exists(core_path):
                    tree = ET.parse(core_path)
                    root = tree.getroot()
                    
                    # Base64 인코딩된 데이터를 메타데이터로 삽입
                    encoded_data = base64.b64encode(data).decode('utf-8')
                    
                    # 데이터를 여러 필드로 분할
                    chunk_size = 200
                    chunks = [encoded_data[i:i+chunk_size] 
                             for i in range(0, len(encoded_data), chunk_size)]
                    
                    # 기존 요소들 수정 또는 새로 추가
                    metadata_fields = [
                        ('title', 'dc:title'),
                        ('subject', 'dc:subject'),
                        ('description', 'dc:description'),
                        ('keywords', 'cp:keywords'),
                        ('category', 'cp:category')
                    ]
                    
                    for i, chunk in enumerate(chunks):
                        if i < len(metadata_fields):
                            field_name, field_tag = metadata_fields[i]
                            
                            # 기존 요소 찾기
                            existing_elem = root.find(field_tag, self.namespaces)
                            if existing_elem is not None:
                                existing_elem.text = chunk
                            else:
                                # 새 요소 추가
                                new_elem = ET.SubElement(root, field_tag)
                                new_elem.text = chunk
                        else:
                            # 사용자 정의 속성으로 추가
                            custom_tag = f"cp:custom{i}"
                            custom_elem = ET.SubElement(root, custom_tag)
                            custom_elem.text = chunk
                    
                    # 헤더 정보 추가
                    header_info = {
                        'chunks': len(chunks),
                        'total_length': len(encoded_data),
                        'checksum': hashlib.md5(data).hexdigest()
                    }
                    
                    header_elem = ET.SubElement(root, 'cp:steganoHeader')
                    header_elem.text = json.dumps(header_info)
                    
                    # XML 저장
                    tree.write(core_path, encoding='utf-8', xml_declaration=True)
                
                # 수정된 DOCX 파일로 압축
                with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as output_zip:
                    for root_dir, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root_dir, file)
                            arc_name = os.path.relpath(file_path, temp_dir)
                            output_zip.write(file_path, arc_name)
                
                return True
                
        except Exception as e:
            self.logger.error(f"메타데이터 임베딩 실패: {str(e)}")
            return False
    
    def _extract_from_metadata(self, docx_path: str) -> bytes:
        """메타데이터에서 데이터 추출"""
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # core.xml 읽기
                with docx_zip.open('docProps/core.xml') as core_file:
                    tree = ET.parse(core_file)
                    root = tree.getroot()
                    
                    # 헤더 정보 찾기
                    header_elem = root.find('cp:steganoHeader', self.namespaces)
                    if header_elem is None or not header_elem.text:
                        raise ValueError("스테가노그래피 헤더를 찾을 수 없습니다")
                    
                    header_info = json.loads(header_elem.text)
                    chunks_count = header_info['chunks']
                    total_length = header_info['total_length']
                    expected_checksum = header_info['checksum']
                    
                    # 데이터 청크들 수집
                    metadata_fields = [
                        'dc:title', 'dc:subject', 'dc:description', 
                        'cp:keywords', 'cp:category'
                    ]
                    
                    encoded_chunks = []
                    for i in range(chunks_count):
                        if i < len(metadata_fields):
                            elem = root.find(metadata_fields[i], self.namespaces)
                            if elem is not None and elem.text:
                                encoded_chunks.append(elem.text)
                        else:
                            # 사용자 정의 속성
                            custom_tag = f'cp:custom{i}'
                            elem = root.find(custom_tag, self.namespaces)
                            if elem is not None and elem.text:
                                encoded_chunks.append(elem.text)
                    
                    # 데이터 복원
                    encoded_data = ''.join(encoded_chunks)[:total_length]
                    decoded_data = base64.b64decode(encoded_data.encode('utf-8'))
                    
                    # 체크섬 검증
                    if hashlib.md5(decoded_data).hexdigest() != expected_checksum:
                        raise ValueError("데이터 무결성 검증 실패")
                    
                    return decoded_data
                    
        except Exception as e:
            raise ValueError(f"메타데이터 추출 실패: {str(e)}")
    
    def _embed_in_whitespace(self, docx: 'Document', binary_data: str) -> bool:
        """공백 문자 패턴으로 데이터 임베딩"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                return False
            
            bit_index = 0
            
            for paragraph in docx.paragraphs:
                if bit_index >= len(binary_data):
                    break
                
                # 문단 끝에 공백 패턴 추가
                if paragraph.text.strip():  # 빈 문단이 아닌 경우만
                    # 비트에 따라 공백 패턴 결정
                    if bit_index < len(binary_data):
                        bit = binary_data[bit_index]
                        if bit == '1':
                            # 비트 1: 탭 + 스페이스
                            paragraph.text += '\t '
                        else:
                            # 비트 0: 스페이스 + 스페이스
                            paragraph.text += '  '
                        bit_index += 1
            
            return bit_index > 0
            
        except Exception as e:
            self.logger.error(f"공백 문자 임베딩 실패: {str(e)}")
            return False
    
    def _embed_in_styles(self, docx: 'Document', binary_data: str) -> bool:
        """텍스트 스타일 미세 조정으로 데이터 임베딩"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                return False
            
            bit_index = 0
            
            for paragraph in docx.paragraphs:
                if bit_index >= len(binary_data):
                    break
                
                for run in paragraph.runs:
                    if bit_index >= len(binary_data):
                        break
                    
                    if run.text.strip():  # 텍스트가 있는 run만
                        bit = binary_data[bit_index]
                        
                        # 비트에 따라 스타일 미세 조정
                        if bit == '1':
                            # 비트 1: 폰트 크기 미세 증가
                            if run.font.size:
                                current_size = run.font.size.pt
                                run.font.size = Pt(current_size + 0.1)
                            else:
                                run.font.size = Pt(12.1)
                        else:
                            # 비트 0: 문자 간격 미세 조정
                            run.font.size = Pt(12.0)
                        
                        bit_index += 1
            
            return bit_index > 0
            
        except Exception as e:
            self.logger.error(f"스타일 임베딩 실패: {str(e)}")
            return False
    
    def _add_hidden_text(self, docx: 'Document', message: str) -> bool:
        """숨김 텍스트 추가"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                return False
            
            # 문서 끝에 숨김 텍스트 추가
            hidden_paragraph = docx.add_paragraph()
            hidden_run = hidden_paragraph.add_run(message)
            
            # 텍스트를 숨김으로 설정
            hidden_run.font.hidden = True
            hidden_run.font.color.rgb = None  # 투명
            
            return True
            
        except Exception as e:
            self.logger.error(f"숨김 텍스트 추가 실패: {str(e)}")
            return False
    
    def embed_message(self, docx_path: str, message: str, output_path: str,
                     password: Optional[str] = None) -> bool:
        """DOCX에 메시지 임베딩"""
        try:
            self.logger.info(f"DOCX 스테가노그래피 임베딩 시작: {docx_path}")
            
            if not PYTHON_DOCX_AVAILABLE:
                raise ImportError("python-docx가 필요합니다. pip install python-docx로 설치하세요.")
            
            # 파일 검증
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"DOCX 파일을 찾을 수 없습니다: {docx_path}")
            
            # 데이터 준비
            if password:
                encrypted_message = self._encrypt_data(message, password)
                data_to_embed = encrypted_message
            else:
                data_to_embed = message.encode('utf-8')
            
            # 헤더 생성
            header = {
                'length': len(data_to_embed),
                'encrypted': password is not None,
                'checksum': hashlib.md5(data_to_embed).hexdigest(),
                'methods': {
                    'metadata': self.metadata_method,
                    'style': self.style_method,
                    'whitespace': self.whitespace_method,
                    'hidden_text': self.hidden_text_method
                }
            }
            header_json = json.dumps(header).encode('utf-8')
            
            # 전체 데이터
            header_length = len(header_json)
            full_data = struct.pack('<I', header_length) + header_json + data_to_embed
            
            # 메타데이터 방식
            if self.metadata_method:
                try:
                    if self._embed_in_metadata(docx_path, full_data, output_path):
                        self.logger.debug("메타데이터 방식으로 임베딩 완료")
                    else:
                        # 메타데이터 방식 실패 시 일반 복사
                        shutil.copy2(docx_path, output_path)
                except Exception as e:
                    self.logger.warning(f"메타데이터 임베딩 실패: {str(e)}")
                    shutil.copy2(docx_path, output_path)
            else:
                shutil.copy2(docx_path, output_path)
            
            # 다른 방식들 (python-docx 사용)
            if self.style_method or self.whitespace_method or self.hidden_text_method:
                try:
                    docx = Document(output_path)
                    binary_data = ''.join(format(byte, '08b') for byte in full_data)
                    
                    # 스타일 방식
                    if self.style_method:
                        if self._embed_in_styles(docx, binary_data):
                            self.logger.debug("스타일 방식으로 임베딩 완료")
                    
                    # 공백 문자 방식
                    if self.whitespace_method:
                        if self._embed_in_whitespace(docx, binary_data):
                            self.logger.debug("공백 문자 방식으로 임베딩 완료")
                    
                    # 숨김 텍스트 방식
                    if self.hidden_text_method:
                        if self._add_hidden_text(docx, message):
                            self.logger.debug("숨김 텍스트 방식으로 임베딩 완료")
                    
                    # 문서 저장
                    docx.save(output_path)
                    
                except Exception as e:
                    self.logger.warning(f"추가 방식 임베딩 실패: {str(e)}")
            
            self.logger.info(f"DOCX 스테가노그래피 임베딩 완료: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"임베딩 중 오류 발생: {str(e)}")
            return False
    
    def extract_message(self, docx_path: str, password: Optional[str] = None) -> str:
        """DOCX에서 메시지 추출"""
        try:
            self.logger.info(f"DOCX 스테가노그래피 추출 시작: {docx_path}")
            
            if not os.path.exists(docx_path):
                raise FileNotFoundError(f"DOCX 파일을 찾을 수 없습니다: {docx_path}")
            
            # 메타데이터에서 추출 시도
            try:
                full_data = self._extract_from_metadata(docx_path)
                
                # 헤더 파싱
                header_length = struct.unpack('<I', full_data[:4])[0]
                header_json = full_data[4:4+header_length]
                header = json.loads(header_json.decode('utf-8'))
                
                # 실제 데이터 추출
                data_bytes = full_data[4+header_length:4+header_length+header['length']]
                
                # 체크섬 검증
                if hashlib.md5(data_bytes).hexdigest() != header['checksum']:
                    raise ValueError("데이터 무결성 검증 실패")
                
                # 복호화
                if header['encrypted']:
                    if not password:
                        raise ValueError("암호화된 데이터이지만 패스워드가 제공되지 않았습니다.")
                    message = self._decrypt_data(data_bytes, password)
                else:
                    message = data_bytes.decode('utf-8')
                
                self.logger.info("DOCX 스테가노그래피 추출 완료 (메타데이터)")
                return message
                
            except Exception as e:
                self.logger.debug(f"메타데이터 추출 실패: {str(e)}")
            
            # 숨김 텍스트에서 추출 시도
            if PYTHON_DOCX_AVAILABLE:
                try:
                    docx = Document(docx_path)
                    
                    # 숨김 텍스트 찾기
                    for paragraph in docx.paragraphs:
                        for run in paragraph.runs:
                            if run.font.hidden:
                                return run.text
                    
                except Exception as e:
                    self.logger.debug(f"숨김 텍스트 추출 실패: {str(e)}")
            
            raise ValueError("숨겨진 메시지를 찾을 수 없습니다.")
            
        except Exception as e:
            self.logger.error(f"추출 중 오류 발생: {str(e)}")
            raise
    
    def get_capacity(self, docx_path: str) -> int:
        """DOCX의 임베딩 용량 계산"""
        try:
            if not os.path.exists(docx_path):
                return 0
            
            capacity = 0
            
            # 메타데이터 용량
            if self.metadata_method:
                capacity += 5 * 200  # 주요 메타데이터 필드들
                capacity += 10 * 200  # 추가 커스텀 필드들
            
            # 스타일 및 공백 방식 용량
            if PYTHON_DOCX_AVAILABLE:
                try:
                    docx = Document(docx_path)
                    
                    # 문단 및 런 수 기반 계산
                    total_runs = 0
                    total_paragraphs = len(docx.paragraphs)
                    
                    for paragraph in docx.paragraphs:
                        total_runs += len(paragraph.runs)
                    
                    if self.style_method:
                        capacity += total_runs  # 런당 1비트
                    
                    if self.whitespace_method:
                        capacity += total_paragraphs  # 문단당 1비트
                    
                    if self.hidden_text_method:
                        capacity += 1000  # 숨김 텍스트로 약 1KB
                    
                except:
                    # Document 로딩 실패 시 기본값
                    capacity += 5000
            
            # 헤더 오버헤드 고려
            overhead = 1024
            return max(0, capacity // 8 - overhead)  # 바이트로 변환
            
        except Exception:
            return 0
    
    def is_suitable_docx(self, docx_path: str) -> Dict[str, Any]:
        """DOCX가 스테가노그래피에 적합한지 분석"""
        try:
            if not os.path.exists(docx_path):
                return {
                    'suitable': False,
                    'reason': '파일을 찾을 수 없습니다',
                    'score': 0.0
                }
            
            file_size = os.path.getsize(docx_path)
            
            analysis = {
                'suitable': False,
                'score': 0.0,
                'file_size_mb': file_size / (1024 * 1024),
                'paragraph_count': 0,
                'run_count': 0,
                'has_images': False,
                'recommendations': []
            }
            
            if PYTHON_DOCX_AVAILABLE:
                try:
                    docx = Document(docx_path)
                    
                    # 문서 구조 분석
                    paragraph_count = len(docx.paragraphs)
                    run_count = sum(len(p.runs) for p in docx.paragraphs)
                    text_length = sum(len(p.text) for p in docx.paragraphs)
                    
                    # 이미지 확인 (간단한 방법)
                    has_images = any('graphicData' in str(p._element.xml) 
                                   for p in docx.paragraphs)
                    
                    # 점수 계산
                    size_score = min(1.0, file_size / (100 * 1024))  # 100KB 기준
                    paragraph_score = min(1.0, paragraph_count / 50.0)  # 50문단 기준
                    run_score = min(1.0, run_count / 200.0)  # 200런 기준
                    text_score = min(1.0, text_length / 2000.0)  # 2000자 기준
                    
                    overall_score = (size_score * 0.2 + paragraph_score * 0.3 + 
                                   run_score * 0.3 + text_score * 0.2)
                    
                    analysis.update({
                        'suitable': overall_score >= 0.4,
                        'score': overall_score,
                        'paragraph_count': paragraph_count,
                        'run_count': run_count,
                        'text_length': text_length,
                        'has_images': has_images,
                        'capacity_bytes': self.get_capacity(docx_path)
                    })
                    
                    # 권장사항
                    if paragraph_count < 10:
                        analysis['recommendations'].append("더 많은 문단이 있는 문서 사용 권장")
                    if run_count < 50:
                        analysis['recommendations'].append("더 복잡한 서식이 있는 문서 사용 권장")
                    if file_size < 50 * 1024:
                        analysis['recommendations'].append("더 큰 문서 사용 권장")
                    
                except Exception as e:
                    analysis.update({
                        'suitable': False,
                        'reason': f'문서 분석 실패: {str(e)}',
                        'score': 0.0
                    })
            else:
                analysis.update({
                    'suitable': False,
                    'reason': 'python-docx 라이브러리가 필요합니다',
                    'score': 0.0
                })
            
            return analysis
            
        except Exception as e:
            return {
                'suitable': False,
                'reason': f'분석 중 오류: {str(e)}',
                'score': 0.0
            }